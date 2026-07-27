"""测试 scripts/book_ingest.py：书籍（PDF/DOCX/TXT）→ 合成逐字稿 → 入库的适配器。

核心承诺：产出的合成逐字稿必须满足 parse.py 的两条硬约束（文件名 14 位日期前缀、
`发言人(HH:MM:SS): 文本` 行格式），这样才能被现有 ingest_new.ingest_pending() 原样
入库，不用改动 parse/chunk/ingest 任何一行。

PDF/DOCX 的底层解析（pypdf/python-docx 本身是否正确读取二进制格式）不是本模块的
职责，测试里 mock 掉；DOCX 用真实的 python-docx round-trip（构造真文件成本很低，
比 mock 更能验证我们真的调对了 API）。
"""
import json
import re
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument

from config import FILENAME_DATETIME_RE, TRANSCRIPT_LINE_RE
from scripts import book_ingest


FNAME_RE = re.compile(FILENAME_DATETIME_RE)
LINE_RE = re.compile(TRANSCRIPT_LINE_RE)


# ── extract_pdf_text ─────────────────────────────────────────────────────


def _fake_pdf_reader(page_texts):
    """构造一个假的 pypdf.PdfReader 替身：.pages 是一串带 extract_text() 的对象。"""
    pages = []
    for t in page_texts:
        page = MagicMock()
        page.extract_text.return_value = t
        pages.append(page)
    reader = MagicMock()
    reader.pages = pages
    return reader


def test_extract_pdf_text_joins_pages(monkeypatch):
    reader = _fake_pdf_reader(["第一页正文" * 10, "第二页正文" * 10])
    monkeypatch.setattr(book_ingest.pypdf, "PdfReader", MagicMock(return_value=reader))

    text = book_ingest.extract_pdf_text(Path("dummy.pdf"))

    assert "第一页正文" in text
    assert "第二页正文" in text
    assert text.count("\n") >= 1  # 页间有分隔


def test_extract_pdf_text_raises_on_scanned_pdf(monkeypatch):
    """页均文字量太低（典型扫描件特征）应该报错提示需要 OCR，而不是静默产出空文档。"""
    reader = _fake_pdf_reader(["", "", ""])
    monkeypatch.setattr(book_ingest.pypdf, "PdfReader", MagicMock(return_value=reader))

    with pytest.raises(ValueError, match="OCR"):
        book_ingest.extract_pdf_text(Path("scanned.pdf"))


def test_extract_pdf_text_empty_pdf_no_pages(monkeypatch):
    """零页 PDF 不应该除零崩溃，应该走扫描件同一条报错路径。"""
    reader = _fake_pdf_reader([])
    monkeypatch.setattr(book_ingest.pypdf, "PdfReader", MagicMock(return_value=reader))

    with pytest.raises(ValueError):
        book_ingest.extract_pdf_text(Path("empty.pdf"))


# ── extract_docx_text ────────────────────────────────────────────────────


def test_extract_docx_text_real_roundtrip(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("第一段：甲子乙丑海中金。")
    doc.add_paragraph("")  # 空段落应被跳过
    doc.add_paragraph("第二段：丙寅丁卯爐中火。")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    text = book_ingest.extract_docx_text(path)

    assert text == "第一段：甲子乙丑海中金。\n第二段：丙寅丁卯爐中火。"


# ── extract_txt_text ─────────────────────────────────────────────────────


def test_extract_txt_text_reads_utf8(tmp_path):
    path = tmp_path / "test.txt"
    path.write_text("五行精紀\n木火土金水", encoding="utf-8")

    assert book_ingest.extract_txt_text(path) == "五行精紀\n木火土金水"


# ── extract_text（分发）──────────────────────────────────────────────────


def test_extract_text_dispatches_by_suffix(tmp_path, monkeypatch):
    called = MagicMock(return_value="来自 txt 提取器")
    monkeypatch.setitem(book_ingest._EXTRACTORS, ".txt", called)
    path = tmp_path / "x.txt"
    path.write_text("占位", encoding="utf-8")

    result = book_ingest.extract_text(path)

    called.assert_called_once_with(path)
    assert result == "来自 txt 提取器"


def test_extract_text_unsupported_suffix_raises(tmp_path):
    path = tmp_path / "book.epub"
    path.write_text("占位", encoding="utf-8")

    with pytest.raises(ValueError, match="不支持"):
        book_ingest.extract_text(path)


# ── clean_book_text ──────────────────────────────────────────────────────


def test_clean_book_text_strips_noise_and_dedup():
    raw = "\n".join([
        "3/25/26, 9:31 AM",             # ctext.org 抓取时间戳（独立一行的情况）
        "三命通會 : 卷一",
        "12",                            # 纯页码
        "在 \"序\" 中檢索：",              # 检索框提示
        "納音之法，同類娶妻。",
        "納音之法，同類娶妻。",           # 连续重复（运行页眉/OCR 重复）应去重
        "",                               # 空行
        "隔八生子，律呂相生之法也。",
    ])

    cleaned = book_ingest.clean_book_text(raw)
    lines = cleaned.splitlines()

    assert "3/25/26, 9:31 AM" not in cleaned
    assert "12" not in lines
    assert "在 \"序\" 中檢索：" not in cleaned
    assert lines.count("納音之法，同類娶妻。") == 1
    assert "三命通會 : 卷一" in cleaned
    assert "隔八生子，律呂相生之法也。" in cleaned


def test_clean_book_text_strips_noise_glued_with_running_title():
    """真实数据里的样子：浏览器"打印为 PDF"页眉是时间戳 + 当页运行标题黏在同一行
    （如 "3/25/26, 9:31 AM 三命通會 : 卷一 - 中國哲學書電子化計劃"），不是独立一行。
    每页标题不同，普通去重/整行匹配都抓不到，必须按前缀丢弃整行。"""
    raw = "\n".join([
        "3/25/26, 9:31 AM 三命通會  : 卷一  - 中國哲學書電子化計劃",
        "納音之法，同類娶妻。",
        "3/25/26, 9:31 AM 三命通會  : 卷二  - 中國哲學書電子化計劃",  # 下一页，标题变了
        "隔八生子，律呂相生之法也。",
    ])

    cleaned = book_ingest.clean_book_text(raw)

    assert "中國哲學書電子化計劃" not in cleaned
    assert "納音之法，同類娶妻。" in cleaned
    assert "隔八生子，律呂相生之法也。" in cleaned


def test_clean_book_text_strips_glued_url_prefix_keeps_real_content():
    """真实数据（三命通會等 ctext.org 来源）每页开头都黏着来源 URL + 页码分数，
    和正文没有换行分隔，必须只挖掉噪音子串、保留正文，不能整行丢弃。"""
    raw = "https://ctext.org/wiki.pl?if=gb&chapter=212352 1/23金主於西，應秋。金之為言禁也。"

    cleaned = book_ingest.clean_book_text(raw)

    assert "ctext.org" not in cleaned
    assert cleaned == "金主於西，應秋。金之為言禁也。"


def test_clean_book_text_strips_glued_urn_suffix_keeps_real_content():
    """每页末尾常黏着 "URNctp:wsNNNNNN喜歡我們的網站？" 引导语，同样和正文黏在一行。"""
    raw = "論太歲論歲運五行精紀宋廖中URN:ctp:ws81208喜歡我們的網站？"

    cleaned = book_ingest.clean_book_text(raw)

    assert "喜歡我們的網站" not in cleaned
    assert "URN:ctp:ws" not in cleaned
    assert cleaned == "論太歲論歲運五行精紀宋廖中"


def test_clean_book_text_drops_standalone_footer_lines():
    raw = "\n".join([
        "正文第一句話。",
        "請支持我們的發展。",
        "網站的設計與内容(c) 版權 2006-2026。如果您想引用本網站上的内容，"
        "請同時加上至本站的鏈接：http://ctext.org/zh。",
        "正文第二句話。",
    ])

    cleaned = book_ingest.clean_book_text(raw)

    assert "請支持我們的發展" not in cleaned
    assert "網站的設計與内容" not in cleaned
    assert "正文第一句話。" in cleaned
    assert "正文第二句話。" in cleaned


def test_clean_book_text_empty_input_returns_empty():
    assert book_ingest.clean_book_text("") == ""
    assert book_ingest.clean_book_text("\n\n   \n") == ""


def test_clean_book_text_char_per_line_preserves_digits_and_repeated_chars():
    """逐字拆行时不能按「整行」去页码/去重，否则会吃掉 URN 数字和叠字。"""
    # 含连续重复字「人人」+ 数字串，全部一行一字；凑满判定阈值
    chars = list("甲子日人人喜土URN81208乙丑丙寅丁卯戊辰己")
    assert len(chars) >= book_ingest._CHAR_PER_LINE_MIN_LINES
    raw = "\n".join(chars)

    cleaned = book_ingest.clean_book_text(raw)
    flat = cleaned.replace("\n", "")

    assert "人人" in flat
    assert "81208" in flat
    assert "URN" in flat


def test_is_char_per_line_text_detects_vertical_pdf_shape():
    assert book_ingest._is_char_per_line_text("\n".join(list("甲" * 30))) is True
    assert book_ingest._is_char_per_line_text("一整段正常正文。\n第二段也正常。") is False
    assert book_ingest._is_char_per_line_text("\n".join(list("短"))) is False  # 样本太短


# ── _reflow_to_paragraphs ────────────────────────────────────────────────


def test_reflow_merges_one_char_per_line_until_sentence_end():
    """竖排 PDF 典型产物：pypdf 把每个字拆成一行。真实数据（滴天髓）证实拼接后
    仍是正确顺序的连贯文本，问题只在断行粒度——这里验证会重新聚合成正常段落，
    而不是保留 1 字一"行"（那样每行都要背 16 字符的说话人/时间戳前缀，噪音占比
    能到 94%）。"""
    chars = list("滴天髓序不見夫氣之塞兩間而彌六合乎沛然而達灼然而炎蕃然而昌毅然而剛厚然")
    text = "\n".join(chars)

    paragraphs = book_ingest._reflow_to_paragraphs(text, target_len=200)

    assert len(paragraphs) == 1  # 全文远小于 target_len，且没有句末标点提前收段
    assert paragraphs[0] == "".join(chars)
    assert all(len(p) <= 200 or p.count("\n") == 0 for p in paragraphs)


def test_reflow_closes_paragraph_on_sentence_end_even_if_short():
    text = "\n".join(["第", "一", "句", "。", "第", "二", "句", "。"])

    paragraphs = book_ingest._reflow_to_paragraphs(text, target_len=200)

    assert paragraphs == ["第一句。", "第二句。"]


def test_reflow_force_closes_when_exceeding_target_len():
    # 200 个「字」且都不含句末标点 → 必须被 target_len 强制切段，不能无限累积
    text = "\n".join(["字"] * 50)  # 50 字符，未达 target_len=10，用小阈值验证强制切段

    paragraphs = book_ingest._reflow_to_paragraphs(text, target_len=10)

    assert all(len(p) <= 10 for p in paragraphs)
    assert "".join(paragraphs) == "字" * 50


def test_reflow_preserves_paragraph_per_line_when_source_already_well_formed():
    """docx/网页抓取来源通常一行一段、且每段已经以句末标点结尾——这种输入应该
    基本保持原有分段粒度不变。"""
    text = "第一段內容說明清楚。\n第二段內容也說明清楚。"

    paragraphs = book_ingest._reflow_to_paragraphs(text, target_len=200)

    assert paragraphs == ["第一段內容說明清楚。", "第二段內容也說明清楚。"]


def test_reflow_empty_text_returns_empty_list():
    assert book_ingest._reflow_to_paragraphs("") == []


# ── convert_book_to_transcript ───────────────────────────────────────────


def test_convert_book_to_transcript_writes_valid_transcript(tmp_path, monkeypatch):
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "第一句話。\n第二句話。")

    out_path = book_ingest.convert_book_to_transcript(
        Path("三命通會.pdf"), workspace_id="bazhai-ziwu", seq=0
    )

    assert out_path.parent == tmp_path
    assert out_path.exists()
    assert FNAME_RE.match(out_path.name), f"文件名不满足 14 位日期前缀: {out_path.name}"
    assert "三命通會" in out_path.name

    lines = out_path.read_text(encoding="utf-8").splitlines()
    assert lines  # 非空
    for line in lines:
        m = LINE_RE.match(line)
        assert m, f"合成逐字稿行不满足格式: {line!r}"
        speaker, ts, body = m.groups()
        assert speaker == "原文"
        assert ts == "00:00:00"
        assert body in ("第一句話。", "第二句話。")


def test_convert_book_to_transcript_seq_makes_filenames_unique(tmp_path, monkeypatch):
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "内容")

    p0 = book_ingest.convert_book_to_transcript(Path("书A.pdf"), "ws", seq=0)
    p1 = book_ingest.convert_book_to_transcript(Path("书B.pdf"), "ws", seq=1)

    assert p0 != p1
    assert p0.name.startswith("19000101000000_")
    assert p1.name.startswith("19000101000001_")


def test_convert_book_to_transcript_raises_on_empty_extracted_text(tmp_path, monkeypatch):
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "   \n   ")

    with pytest.raises(ValueError, match="没有可用文本"):
        book_ingest.convert_book_to_transcript(Path("空书.pdf"), "ws", seq=0)


def test_convert_book_to_transcript_does_not_bloat_on_char_per_line_source(tmp_path, monkeypatch):
    """回归测试：extract_text 返回逐字一行的文本（竖排 PDF 典型产物）时，产出的合成
    逐字稿不应该给每个字都套一条完整的 "原文(00:00:00): " 前缀——那样前缀噪音字符数
    会远超正文字符数。"""
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    chars = list("滴天髓序不見夫氣之塞兩間而彌六合乎沛然而達灼然而炎")
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "\n".join(chars))

    out_path = book_ingest.convert_book_to_transcript(Path("滴天髓.pdf"), "ws", seq=0)
    content = out_path.read_text(encoding="utf-8")
    n_lines = len(content.splitlines())
    prefix_chars = n_lines * len(f"{book_ingest._BOOK_SPEAKER}({book_ingest._BOOK_TS}): ")
    body_chars = len("".join(chars))

    assert n_lines <= 2, f"逐字一行的输入应该被重新聚合成极少数段落，实际 {n_lines} 行"
    assert prefix_chars < body_chars, "前缀噪音字符数不应该超过正文字符数"


def test_strip_glued_noise_removes_prefix_and_suffix():
    """_strip_glued_noise 是 reflow 后的二次清洗：挖掉黏在段落首尾的 ctext 噪音子串。"""
    text = (
        "https://ctext.org/wiki.pl?if=gb&chapter=1 1/23"
        "正文内容URN:ctp:ws81208喜歡我們的網站？"
    )
    cleaned = book_ingest._strip_glued_noise(text)
    assert cleaned == "正文内容"
    assert "ctext.org" not in cleaned
    assert "URN:ctp" not in cleaned
    assert "喜歡我們的網站" not in cleaned


def test_strip_glued_noise_url_without_page_fraction_keeps_chinese():
    """合并 txt 常见形态：URL 后直接接中文标题，无 '1/23' 页码分数。"""
    text = "來源網址: https://ctext.org/wiki.pl?if=gb&chapter=81208五行精紀 : 序"
    cleaned = book_ingest._strip_glued_noise(text)
    assert "ctext.org" not in cleaned
    assert "來源網址" not in cleaned
    assert "五行精紀 : 序" in cleaned


def test_strip_glued_noise_spaced_urn_and_footer():
    """部分 PDF 把 URN 拆成字母间带空格，且后面黏赞助/版权长文。"""
    text = (
        "正文R N :  c t p : w s 5 5 0 6 5 6喜歡我們的網站？請⽀持我們的發展。"
        "站的鏈接：http://ctext.org/zh。請注意：嚴禁使⽤自動下載軟体下載本網站的大量網頁，"
        "違者自動封鎖，不另行通知。"
    )
    cleaned = book_ingest._strip_glued_noise(text)
    assert "ctext.org" not in cleaned
    assert "喜歡我們的網站" not in cleaned
    assert "URN" not in cleaned.upper().replace(" ", "") or "ctp" not in cleaned
    assert "正文" in cleaned


def test_clean_book_text_drops_ctext_section_and_copyright_paragraphs():
    """合并 txt 每篇末尾的 (c)版权 / 引用说明 / ICP 应整段丢弃。"""
    raw = "\n".join([
        "納音之法，同類娶妻。",
        "(c)版權2006-2026。",
        "如果您想引用本網站上的内容，請同時加上至本站的鏈接：http://ctext.org/zh。"
        "請注意：嚴禁使用自動下載軟体下載本網站的大量網頁，違者自動封鎖，不另行通知。",
        "沪ICP备09015720号-3若有任何意見或建議，請在此提出。",
        "隔八生子，律呂相生之法也。",
    ])
    cleaned = book_ingest.clean_book_text(raw)
    assert "版權" not in cleaned
    assert "ctext.org" not in cleaned
    assert "ICP" not in cleaned
    assert "納音之法，同類娶妻。" in cleaned
    assert "隔八生子，律呂相生之法也。" in cleaned


def test_convert_book_to_transcript_strips_char_per_line_glued_noise(tmp_path, monkeypatch):
    """回归：噪音本身也是逐字拆行时，clean_book_text 按单行匹配抓不到；
    必须在 reflow 拼回连续子串后再跑 _strip_glued_noise，否则检索会命中
    'URN:ctp' / '喜歡我們的網站' 这类 ctext.org 页脚。"""
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    chars = list("正文正文URN:ctp:ws81208喜歡我們的網站？正文正文")
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "\n".join(chars))

    out_path = book_ingest.convert_book_to_transcript(Path("noisy.pdf"), "ws", seq=0)
    content = out_path.read_text(encoding="utf-8")

    assert "URN:ctp" not in content
    assert "喜歡我們的網站" not in content
    assert "正文正文" in content


def test_convert_book_to_transcript_strips_char_per_line_ctext_url_prefix(tmp_path, monkeypatch):
    """回归：ctext.org URL 前缀同样可能逐字拆行，reflow 后应被挖掉并保留正文。"""
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    chars = list("https://ctext.org/wiki.pl?if=gb&chapter=212352 1/23金主於西，應秋。")
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "\n".join(chars))

    out_path = book_ingest.convert_book_to_transcript(Path("三命.pdf"), "ws", seq=0)
    content = out_path.read_text(encoding="utf-8")

    assert "ctext.org" not in content
    assert "金主於西，應秋。" in content


def test_convert_book_to_transcript_raises_when_only_noise_after_reflow(tmp_path, monkeypatch):
    """整份提取文本 reflow 后全是噪音（无正文）时，应报「没有可用文本」。"""
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    # 逐字拆开的纯噪音：reflow 后会被 _strip_glued_noise / _NOISE_LINE_RE 清空
    chars = list("URN:ctp:ws81208喜歡我們的網站？")
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "\n".join(chars))

    with pytest.raises(ValueError, match="没有可用文本"):
        book_ingest.convert_book_to_transcript(Path("纯噪音.pdf"), "ws", seq=0)


def test_convert_book_to_transcript_custom_title(tmp_path, monkeypatch):
    monkeypatch.setattr(book_ingest, "RAW_DIR", lambda ws=None: tmp_path)
    monkeypatch.setattr(book_ingest, "extract_text", lambda p: "内容")

    out_path = book_ingest.convert_book_to_transcript(
        Path("NLC416-13jh000625-42999_河洛理數.pdf"), "ws", seq=0, title="河洛理數"
    )

    assert out_path.name == "19000101000000_河洛理數.txt"


# ── ingest_book_folder（编排）─────────────────────────────────────────────


def test_ingest_book_folder_converts_then_ingests(tmp_path, monkeypatch):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    (src_dir / "a.pdf").write_bytes(b"fake-pdf")
    (src_dir / "b.docx").write_bytes(b"fake-docx")
    (src_dir / "notes.epub").write_bytes(b"unsupported")  # 应被忽略，不进 converted/failed

    converted_calls = []

    def fake_convert(path, workspace_id, seq=0, title=None):
        converted_calls.append((path.name, workspace_id, seq))
        out = tmp_path / f"out-{path.name}.txt"
        out.write_text("占位", encoding="utf-8")
        return out

    fake_ingest_only = MagicMock(return_value=None)

    monkeypatch.setattr(book_ingest, "convert_book_to_transcript", fake_convert)
    monkeypatch.setattr(book_ingest, "_ingest_chunks_only", fake_ingest_only)

    result = book_ingest.ingest_book_folder(src_dir, workspace_id="bazhai-ziwu")

    assert sorted(result["converted"]) == ["a.pdf", "b.docx"]
    assert result["convert_failed"] == []
    assert sorted(result["ingested"]) == ["a.pdf", "b.docx"]
    assert result["failed"] == []
    assert fake_ingest_only.call_count == 2
    # 两份都应传了各自唯一的 seq
    seqs = sorted(c[2] for c in converted_calls)
    assert seqs == [0, 1]


def test_ingest_book_folder_one_convert_failure_does_not_block_others(tmp_path, monkeypatch):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    (src_dir / "good.txt").write_bytes(b"ok")
    (src_dir / "scanned.pdf").write_bytes(b"scanned")

    def fake_convert(path, workspace_id, seq=0, title=None):
        if path.name == "scanned.pdf":
            raise ValueError(f"{path.name}: 疑似扫描件，需要先做 OCR")
        out = tmp_path / f"out-{path.name}.txt"
        out.write_text("占位", encoding="utf-8")
        return out

    monkeypatch.setattr(book_ingest, "convert_book_to_transcript", fake_convert)
    monkeypatch.setattr(book_ingest, "_ingest_chunks_only", MagicMock(return_value=None))

    result = book_ingest.ingest_book_folder(src_dir, workspace_id="ws")

    assert result["converted"] == ["good.txt"]
    assert len(result["convert_failed"]) == 1
    assert result["convert_failed"][0]["file"] == "scanned.pdf"
    assert "OCR" in result["convert_failed"][0]["error"]
    assert result["ingested"] == ["good.txt"]


def test_ingest_book_folder_one_ingest_failure_does_not_block_others(tmp_path, monkeypatch):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    (src_dir / "a.txt").write_bytes(b"ok")
    (src_dir / "b.txt").write_bytes(b"ok")

    def fake_convert(path, workspace_id, seq=0, title=None):
        out = tmp_path / f"out-{path.name}.txt"
        out.write_text("占位", encoding="utf-8")
        return out

    def fake_ingest_only(out_path, workspace_id):
        if "a.txt" in out_path.name:
            raise RuntimeError("LanceDB 写入失败")

    monkeypatch.setattr(book_ingest, "convert_book_to_transcript", fake_convert)
    monkeypatch.setattr(book_ingest, "_ingest_chunks_only", fake_ingest_only)

    result = book_ingest.ingest_book_folder(src_dir, workspace_id="ws")

    assert result["converted"] == ["a.txt", "b.txt"]
    assert result["ingested"] == ["b.txt"]
    assert len(result["failed"]) == 1
    assert result["failed"][0]["file"] == "a.txt"
    assert "LanceDB" in result["failed"][0]["error"]


def test_ingest_book_folder_respects_skip_files(tmp_path, monkeypatch):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    (src_dir / "keep.txt").write_bytes(b"ok")
    (src_dir / "skip.pdf").write_bytes(b"ok")

    fake_convert = MagicMock(side_effect=lambda path, workspace_id, seq=0, title=None: tmp_path / "out.txt")
    monkeypatch.setattr(book_ingest, "convert_book_to_transcript", fake_convert)
    monkeypatch.setattr(book_ingest, "_ingest_chunks_only", MagicMock(return_value=None))

    book_ingest.ingest_book_folder(src_dir, workspace_id="ws", skip_files=["skip.pdf"])

    called_names = [c.args[0].name for c in fake_convert.call_args_list]
    assert called_names == ["keep.txt"]


# ── _ingest_chunks_only ───────────────────────────────────────────────────


def test_ingest_chunks_only_writes_chunks_and_calls_ingest(tmp_path, monkeypatch):
    """验证跳过摘要/长期记忆——那是心理咨询专用 prompt（scripts/summarize.py 的
    SYSTEM_INSTRUCTION 硬编码"你是一位心理咨询记录整理助手"，schema 要求
    emotional_tone/psychological_themes），对书籍类内容answer 没有意义，而且会把
    整本书原文一次性砸进单次 LLM 调用（本项目古籍单份 20 万+字符）。这里只做
    parse→chunk→向量化入库，不碰 summarize_session / update_memory。"""
    from dataclasses import dataclass

    @dataclass
    class _FakeChunk:
        id: str
        source_file: str = "book.txt"

    fake_session = MagicMock(source_file="19000101000000_book.txt", session_date="1900-01-01")
    chunks_path = tmp_path / "chunks.jsonl"

    mock_parse = MagicMock(return_value=fake_session)
    mock_chunk_session = MagicMock(return_value=[_FakeChunk("c0"), _FakeChunk("c1")])
    mock_ingest = MagicMock()
    mock_record = MagicMock()

    monkeypatch.setattr(book_ingest, "parse_transcript", mock_parse)
    monkeypatch.setattr(book_ingest, "chunk_session", mock_chunk_session)
    monkeypatch.setattr(book_ingest, "ingest", mock_ingest)
    monkeypatch.setattr(book_ingest, "append_change_record", mock_record)
    monkeypatch.setattr(book_ingest, "CHUNKS_JSONL_PATH", lambda ws=None: chunks_path)

    book_ingest._ingest_chunks_only(Path("19000101000000_book.txt"), workspace_id="bazhai-ziwu")

    mock_chunk_session.assert_called_once()
    mock_ingest.assert_called_once()
    assert mock_ingest.call_args.kwargs["mode"] == "append"
    assert mock_ingest.call_args.kwargs["workspace_id"] == "bazhai-ziwu"
    assert chunks_path.exists()
    written = [json.loads(l) for l in chunks_path.read_text(encoding="utf-8").splitlines()]
    assert [c["id"] for c in written] == ["c0", "c1"]
    mock_record.assert_called_once()
    assert mock_record.call_args.args[0] == "added"
    # 核心断言（结构性，不是运行时 mock）：book_ingest 模块根本不 import summarize_session /
    # update_memory —— 从代码层面保证不会对书籍内容跑心理咨询专用的摘要 prompt。
    assert not hasattr(book_ingest, "summarize_session")
    assert not hasattr(book_ingest, "update_memory")


def test_ingest_chunks_only_is_idempotent_skips_already_indexed(tmp_path, monkeypatch):
    """已入库过的文件重跑应该跳过（不重复 ingest()、不重复写 chunks.jsonl），
    和 ingest_new_file() 的幂等承诺一致——避免同一批书反复跑触发重复向量。"""
    chunks_path = tmp_path / "chunks.jsonl"
    chunks_path.write_text(
        json.dumps({"id": "c0", "source_file": "19000101000000_book.txt"}, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )

    fake_session = MagicMock(source_file="19000101000000_book.txt", session_date="1900-01-01")
    mock_chunk_session = MagicMock()
    mock_ingest = MagicMock()
    mock_record = MagicMock()

    monkeypatch.setattr(book_ingest, "parse_transcript", MagicMock(return_value=fake_session))
    monkeypatch.setattr(book_ingest, "chunk_session", mock_chunk_session)
    monkeypatch.setattr(book_ingest, "ingest", mock_ingest)
    monkeypatch.setattr(book_ingest, "append_change_record", mock_record)
    monkeypatch.setattr(book_ingest, "CHUNKS_JSONL_PATH", lambda ws=None: chunks_path)

    book_ingest._ingest_chunks_only(Path("19000101000000_book.txt"), workspace_id="ws")

    mock_chunk_session.assert_not_called()
    mock_ingest.assert_not_called()
    mock_record.assert_called_once()
    assert mock_record.call_args.args[0] == "skipped"
